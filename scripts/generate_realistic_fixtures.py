#!/usr/bin/env python3
"""Generate synthetic client-style visa evidence fixtures.

The files contain no real personal data. Text PDFs are intentionally simple so
the local extractor can read them without OCR. Scanned PDFs and PNGs include an
`.ocr.txt` sidecar that stands in for Baidu OCR during the offline demo.
"""
from pathlib import Path
import textwrap
import time
import zipfile

import yaml
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "fixtures" / "realistic-materials"


def pdf_escape(value):
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def write_text_pdf(path, title, lines):
    commands = [
        "0.08 0.20 0.36 rg 0 710 612 82 re f",
        "1 1 1 rg BT /F1 17 Tf 54 754 Td (%s) Tj ET" % pdf_escape(title),
        "1 1 1 rg BT /F1 8 Tf 54 730 Td (SYNTHETIC CUSTOMER DOCUMENT) Tj ET",
        "0.94 0.96 0.98 rg 48 610 516 75 re f",
        "0.72 0.76 0.82 RG 48 610 516 75 re S",
    ]
    y = 660
    for index, line in enumerate(lines):
        if index == 5:
            commands.extend([
                "0.12 0.28 0.46 rg 48 %d 516 24 re f" % (y - 6),
                "1 1 1 rg BT /F1 9 Tf 58 %d Td (DOCUMENT DETAILS) Tj ET" % y,
            ])
            y -= 31
        shade = "0.97 0.98 0.99" if index % 2 else "1 1 1"
        commands.extend([
            "%s rg 48 %d 516 20 re f" % (shade, y - 6),
            "0.20 0.24 0.30 rg BT /F1 9 Tf 58 %d Td (%s) Tj ET" % (y, pdf_escape(line)),
            "0.88 0.90 0.93 RG 48 %d m 564 %d l S" % (y - 8, y - 8),
        ])
        y -= 22
    commands.extend([
        "0.70 0.12 0.12 rg BT /F1 8 Tf 54 45 Td (SYNTHETIC TEST DOCUMENT - NOT VALID FOR ANY APPLICATION) Tj ET",
        "0.45 0.48 0.52 rg BT /F1 8 Tf 500 45 Td (Page 1 of 1) Tj ET",
    ])
    stream = "\n".join(commands).encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
         b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
    ]
    data = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, obj in enumerate(objects, 1):
        offsets.append(len(data))
        data.extend(("%d 0 obj\n" % index).encode("ascii"))
        data.extend(obj)
        data.extend(b"\nendobj\n")
    xref = len(data)
    data.extend(("xref\n0 %d\n" % (len(objects) + 1)).encode("ascii"))
    data.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        data.extend(("%010d 00000 n \n" % offset).encode("ascii"))
    data.extend(("trailer << /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n"
                 % (len(objects) + 1, xref)).encode("ascii"))
    path.write_bytes(data)


def write_docx(path, title, paragraphs):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SYNTHETIC VISA EVIDENCE • CASE DEMO-2026-001")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 104, 120)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(24, 52, 86)

    date_line = document.add_paragraph("22 August 2026")
    date_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_line.runs[0].font.size = Pt(9)

    field_lines = [p for p in paragraphs if ":" in p and not p.lower().startswith("i ")]
    prose_lines = [p for p in paragraphs if p not in field_lines]
    if field_lines:
        table = document.add_table(rows=0, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for index, value in enumerate(field_lines):
            cell = table.add_row().cells[0]
            cell.text = value
            cell.vertical_alignment = 1
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EEF3F8" if index % 2 == 0 else "FFFFFF")
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
    for value in prose_lines:
        paragraph = document.add_paragraph(value)
        paragraph.paragraph_format.space_after = Pt(8)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10.5)

    document.add_paragraph("Yours faithfully,")
    signature = document.add_paragraph("____________________________\nAuthorised signatory")
    signature.runs[0].italic = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SYNTHETIC TEST DOCUMENT — NOT VALID FOR IMMIGRATION OR IDENTITY USE")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(170, 40, 40)
    document.core_properties.created = __import__("datetime").datetime(2026, 8, 22)
    document.core_properties.modified = __import__("datetime").datetime(2026, 8, 22)
    document.save(path)
    _normalise_docx_archive(path)


def _normalise_docx_archive(path):
    """Remove wall-clock ZIP metadata so regenerated fixtures are byte-stable."""
    temp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source:
        members = [(item.filename, source.read(item.filename)) for item in source.infolist()]
    with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
        for name, content in members:
            item = zipfile.ZipInfo(name, date_time=(2026, 8, 22, 0, 0, 0))
            item.compress_type = zipfile.ZIP_DEFLATED
            target.writestr(item, content)
    temp.replace(path)


def font(size):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def render_scan(path, title, lines, image_format):
    image = Image.new("RGB", (1240, 1754), "#ebeef2")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((55, 55, 1185, 1699), radius=24, fill="#fbfaf6",
                           outline="#66778a", width=4)
    draw.rectangle((55, 55, 1185, 220), fill="#183b64")
    draw.text((100, 92), title, font=font(43), fill="white")
    draw.text((102, 160), "OFFICIAL-STYLE SYNTHETIC FIXTURE", font=font(20), fill="#cbd8e6")

    is_card = "PASSPORT" in title or "STATUS" in title
    if is_card:
        draw.rounded_rectangle((90, 280, 410, 730), radius=18, fill="#d6dde5",
                               outline="#7f8b99", width=3)
        draw.ellipse((178, 350, 322, 494), fill="#98a6b5")
        draw.rounded_rectangle((145, 500, 355, 680), radius=70, fill="#98a6b5")
        draw.text((142, 690), "PORTRAIT", font=font(18), fill="#5d6976")
        x, y = 465, 300
    else:
        draw.rounded_rectangle((90, 275, 1150, 380), radius=12, fill="#eef3f7",
                               outline="#a5b1be", width=2)
        draw.text((120, 302), "Document reference: DEMO-2026-001",
                  font=font(24), fill="#35485c")
        x, y = 115, 430

    for line in lines:
        wrapped = textwrap.wrap(line, width=46 if is_card else 68) or [""]
        for part in wrapped:
            draw.text((x, y), part, font=font(27), fill="#243241")
            y += 41
        draw.line((x, y + 4, 1125, y + 4), fill="#d3d9df", width=1)
        y += 24
    if "PASSPORT" in title:
        values = {}
        for line in lines:
            if ":" in line:
                key, value = line.split(":", 1)
                values[key.strip().lower()] = value.strip()
        passport_number = values.get("passport number", "EK0000000")
        holder = values.get("holder name", "SYNTHETIC HOLDER").upper().replace(" ", "<")
        expiry = values.get("expiry date", "2029-04-30").replace("-", "")[2:]
        draw.rectangle((90, 1280, 1150, 1510), fill="#e7ebdf", outline="#88937d", width=2)
        draw.text((112, 1330), ("P<CHN%s<<<<<<<<<<<<<<<<<<<<<<<<" % holder)[:42],
                  font=font(25), fill="#263328")
        draw.text((112, 1395), ("%s<CHN9001018F%s<<<<<<<<<<<<<<<<" %
                  (passport_number, expiry))[:42], font=font(25), fill="#263328")
    draw.text((100, 1605), "SYNTHETIC TEST DOCUMENT — NOT VALID FOR TRAVEL OR IDENTITY",
              font=font(22), fill="#a03636")
    if image_format == "PDF":
        fixed_time = time.gmtime(1787356800)  # 2026-08-22T00:00:00Z
        image.save(path, "PDF", resolution=150.0,
                   creationDate=fixed_time, modDate=fixed_time)
    else:
        image.save(path, image_format)


def add_ocr_sidecar(path, lines):
    Path(str(path) + ".ocr.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")


def entry(evidence_id, relative_path, content_type, fields, failures,
          case="pass", note=None, sidecar=False, pair=None):
    result = {
        "case": case,
        "evidence_id": evidence_id,
        "filename": relative_path,
        "content_type": content_type,
        "expected_fields": fields,
        "expected_blocking_failures": failures,
    }
    if note:
        result["scenario"] = note
    if sidecar:
        result["ocr_sidecar"] = relative_path + ".ocr.txt"
    if pair:
        result["pair"] = pair
    return result


def generate():
    for folder in (OUT / "pass", OUT / "fail", OUT / "cross-document"):
        folder.mkdir(parents=True, exist_ok=True)

    fixtures = []

    # 1. Passport: scanned PDF, OCR sidecar.
    lines = ["Holder Name: Mei Ling Chen", "Passport Number: EK1234567",
             "Nationality: Chinese", "Expiry Date: 2029-04-30",
             "Prior Compliant Travel: yes"]
    path = OUT / "pass" / "passport-pass-scanned.pdf"
    render_scan(path, "PASSPORT BIOGRAPHIC PAGE", lines, "PDF")
    add_ocr_sidecar(path, lines)
    fixtures.append(entry("passport", "pass/" + path.name, "application/pdf",
                          {"holder_name": "Mei Ling Chen", "passport_number": "EK1234567",
                           "nationality": "Chinese", "expiry_date": "2029-04-30",
                           "prior_compliant_travel": True}, [], sidecar=True,
                          note="Readable scanned passport with validity beyond the trip."))
    lines = ["Holder Name: Mei Ling Chen", "Passport Number: EK7654321",
             "Nationality: Chinese", "Expiry Date: 2026-12-01",
             "Prior Compliant Travel: yes"]
    path = OUT / "fail" / "passport-fail-expired-scanned.pdf"
    render_scan(path, "PASSPORT BIOGRAPHIC PAGE", lines, "PDF")
    add_ocr_sidecar(path, lines)
    fixtures.append(entry("passport", "fail/" + path.name, "application/pdf",
                          {"holder_name": "Mei Ling Chen", "passport_number": "EK7654321",
                           "nationality": "Chinese", "expiry_date": "2026-12-01",
                           "prior_compliant_travel": True}, ["date_after"], case="fail",
                          sidecar=True, note="Passport expires before the planned return date."))

    # 2. Bank statements: native text PDF.
    lines = ["Account Holder Name: Mei Ling Chen", "Period Start: 2026-02-10",
             "Period End: 2026-08-18", "Closing Balance: 5100.00", "Currency: GBP",
             "Statement number: 08/2026", "Account type: Everyday Current Account",
             "2026-08-18 | CARD PAYMENT | GROCERIES | -82.40 | 5100.00",
             "2026-08-15 | CLIENT PAYMENT | DESIGN RETAINER | +1850.00 | 5182.40",
             "2026-08-09 | BANK TRANSFER | RENT | -920.00 | 3332.40",
             "2026-08-01 | OPENING BALANCE | | | 4252.40"]
    path = OUT / "pass" / "bank_statements-pass.pdf"
    write_text_pdf(path, "PERSONAL CURRENT ACCOUNT STATEMENT", lines)
    fixtures.append(entry("bank_statements", "pass/" + path.name, "application/pdf",
                          {"account_holder_name": "Mei Ling Chen", "period_start": "2026-02-10",
                           "period_end": "2026-08-18", "closing_balance": "5100.00",
                           "currency": "GBP"}, [], note="Recent statement covering over 180 days."))
    lines[0] = "Account Holder Name: Wei Zhang"
    path = OUT / "fail" / "bank_statements-fail-wrong-name.pdf"
    write_text_pdf(path, "PERSONAL CURRENT ACCOUNT STATEMENT", lines)
    fixtures.append(entry("bank_statements", "fail/" + path.name, "application/pdf",
                          {"account_holder_name": "Wei Zhang", "period_start": "2026-02-10",
                           "period_end": "2026-08-18", "closing_balance": "5100.00",
                           "currency": "GBP"}, ["name_matches"], case="fail",
                          note="Statement belongs to a different person."))

    # 3. Itinerary: native text PDF.
    lines = ["Booking Reference: SYNTH-UK-2601", "Passenger Name: Mei Ling Chen",
             "Outbound Date: 2026-10-05", "Return Date: 2027-01-03",
             "Ticket number: 999-2601000001", "Cabin: Economy Flexible",
             "OUTBOUND | ZX201 | Shanghai PVG 09:20 | London LHR 15:35",
             "RETURN | ZX202 | London LHR 18:10 | Shanghai PVG 12:25 +1",
             "Baggage: 1 x 23kg checked; 1 x cabin bag",
             "Status: Reservation held for visa documentation only"]
    path = OUT / "pass" / "travel_itinerary-pass.pdf"
    write_text_pdf(path, "RETURN FLIGHT ITINERARY", lines)
    fixtures.append(entry("travel_itinerary", "pass/" + path.name, "application/pdf",
                          {"passenger_name": "Mei Ling Chen", "outbound_date": "2026-10-05",
                           "return_date": "2027-01-03"}, [], note="Return itinerary matches stated dates."))
    lines[1] = "Passenger Name: Wei Zhang"
    path = OUT / "fail" / "travel_itinerary-fail-wrong-name.pdf"
    write_text_pdf(path, "RETURN FLIGHT ITINERARY", lines)
    fixtures.append(entry("travel_itinerary", "fail/" + path.name, "application/pdf",
                          {"passenger_name": "Wei Zhang", "outbound_date": "2026-10-05",
                           "return_date": "2027-01-03"}, ["name_matches"], case="fail",
                          note="Itinerary names a different passenger."))

    # 4. Accommodation: Word documents.
    lines = ["To: UK Visas and Immigration", "Property reference: SYNTH-M20-1431",
             "Address: 14 Bramhall Road, Manchester M20 3QT", "Host Name: Hui Chen",
             "Stay Start: 2026-10-05", "Stay End: 2027-01-03",
             "Dear Sir or Madam,",
             "I confirm that Mei Ling Chen may stay in the spare bedroom at my home for the full visit.",
             "The property has two bedrooms and no other temporary guests are expected during this period."]
    path = OUT / "pass" / "accommodation_proof-pass.docx"
    write_docx(path, "Accommodation confirmation", lines)
    fixtures.append(entry("accommodation_proof", "pass/" + path.name,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          {"address": "14 Bramhall Road, Manchester M20 3QT", "host_name": "Hui Chen",
                           "stay_start": "2026-10-05", "stay_end": "2027-01-03"}, [],
                          note="Host confirms address and stay dates."))
    lines = ["To: UK Visas and Immigration", "Host Name: Hui Chen",
             "Stay Start: 2026-10-05", "Stay End: 2027-01-03", "Dear Sir or Madam,",
             "I confirm the guest will stay at my home; the property address is omitted from this letter."]
    path = OUT / "fail" / "accommodation_proof-fail-missing-address.docx"
    write_docx(path, "Accommodation confirmation", lines)
    fixtures.append(entry("accommodation_proof", "fail/" + path.name,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          {"host_name": "Hui Chen", "stay_start": "2026-10-05",
                           "stay_end": "2027-01-03"}, ["field_present"], case="fail",
                          note="The document never states the accommodation address."))

    # 5. Invitation letter: Word documents.
    lines = ["To: Entry Clearance Officer", "Applicant: Mei Ling Chen",
             "Sponsor Name: Hui Chen", "Sponsor Address: 14 Bramhall Road, Manchester M20 3QT",
             "Relationship: sister", "Stay Start: 2026-10-05", "Stay End: 2027-01-03",
             "Funding Offered: accommodation only",
             "Dear Entry Clearance Officer,",
             "I invite my sister Mei Ling Chen to visit me in Manchester for a family visit.",
             "She will stay with me at the address above. She will pay her travel and living costs; I will provide accommodation.",
             "I enclose synthetic proof of my UK immigration status for consistency testing."]
    path = OUT / "pass" / "sponsor_invitation_letter-pass.docx"
    write_docx(path, "Invitation letter", lines)
    fixtures.append(entry("sponsor_invitation_letter", "pass/" + path.name,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          {"sponsor_name": "Hui Chen",
                           "sponsor_address": "14 Bramhall Road, Manchester M20 3QT",
                           "relationship": "sister", "stay_start": "2026-10-05",
                           "stay_end": "2027-01-03", "funding_offered": "accommodation only"}, [],
                          note="Invitation identifies host, relationship and arrangements."))
    lines = [line for line in lines if not line.startswith("Relationship:")]
    path = OUT / "fail" / "sponsor_invitation_letter-fail-missing-relationship.docx"
    write_docx(path, "Invitation letter", lines)
    fixtures.append(entry("sponsor_invitation_letter", "fail/" + path.name,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          {"sponsor_name": "Hui Chen",
                           "sponsor_address": "14 Bramhall Road, Manchester M20 3QT",
                           "stay_start": "2026-10-05", "stay_end": "2027-01-03",
                           "funding_offered": "accommodation only"}, ["field_present"], case="fail",
                          note="Invitation does not state the relationship."))

    # 6. Sponsor status: photographed card/image, OCR sidecar.
    lines = ["Sponsor Name: Hui Chen", "Status Type: Indefinite Leave to Remain",
             "Document reference: SYNTH-ILR-8842"]
    path = OUT / "pass" / "sponsor_status_proof-pass.png"
    render_scan(path, "UK STATUS EVIDENCE", lines, "PNG")
    add_ocr_sidecar(path, lines)
    fixtures.append(entry("sponsor_status_proof", "pass/" + path.name, "image/png",
                          {"sponsor_name": "Hui Chen", "status_type": "Indefinite Leave to Remain"},
                          [], sidecar=True, note="Photographed status evidence with readable OCR."))
    lines = ["Sponsor Name: Hui Chen", "Document reference: SYNTH-ILR-8842"]
    path = OUT / "fail" / "sponsor_status_proof-fail-missing-status.png"
    render_scan(path, "UK STATUS EVIDENCE", lines, "PNG")
    add_ocr_sidecar(path, lines)
    fixtures.append(entry("sponsor_status_proof", "fail/" + path.name, "image/png",
                          {"sponsor_name": "Hui Chen"}, ["field_present"], case="fail",
                          sidecar=True, note="Image does not expose the sponsor's status type."))

    # 7. Self-employment evidence: native text PDF bundle summary.
    lines = ["Bundle Reference: SYNTH-BIZ-2026-08", "Business Name: Chen Design Studio",
             "Registration ID: 91310115MA1K3",
             "Tax Year: 2025", "Declared Income: 38000",
             "Business Statement Period Start: 2026-02-10",
             "Business Statement Period End: 2026-08-18",
             "Supporting page 1: business registration extract",
             "Supporting page 2: annual tax filing acknowledgement",
             "Supporting page 3: six-month business account summary",
             "Business status: active", "Registered activity: product and graphic design"]
    path = OUT / "pass" / "self_employment_evidence-pass.pdf"
    write_text_pdf(path, "SELF-EMPLOYMENT EVIDENCE BUNDLE", lines)
    fixtures.append(entry("self_employment_evidence", "pass/" + path.name, "application/pdf",
                          {"business_name": "Chen Design Studio", "registration_id": "91310115MA1K3",
                           "tax_year": "2025", "declared_income": "38000",
                           "business_statement_period_start": "2026-02-10",
                           "business_statement_period_end": "2026-08-18"}, [],
                          note="Registration, tax and business statement facts are all present."))
    lines = ["Business Name: Chen Design Studio", "Registration ID: 91310115MA1K3",
             "Declared Income: 38000", "Business Statement Period Start: 2026-02-10"]
    path = OUT / "fail" / "self_employment_evidence-fail-incomplete.pdf"
    write_text_pdf(path, "SELF-EMPLOYMENT EVIDENCE BUNDLE", lines)
    fixtures.append(entry("self_employment_evidence", "fail/" + path.name, "application/pdf",
                          {"business_name": "Chen Design Studio", "registration_id": "91310115MA1K3",
                           "declared_income": "38000", "business_statement_period_start": "2026-02-10"},
                          ["field_present", "field_present"], case="fail",
                          note="Tax year and business statement end date are missing."))

    # Alternative work branch. Employed clients need this instead of the
    # self-employment bundle.
    lines = ["To: Entry Clearance Officer", "Employee Name: Mei Ling Chen",
             "Employee ID: SHD-1048", "Employer Name: Shanghai Horizon Design Co., Ltd.",
             "Job Title: Senior Product Designer", "Leave Start: 2026-10-02",
             "Leave End: 2027-01-05", "Annual Salary: CNY 420000",
             "Dear Entry Clearance Officer,",
             "This letter confirms that Mei Ling Chen has worked with us full-time since 2021-03-15.",
             "Her paid leave is approved and she is expected to resume work on 2027-01-06.",
             "Human Resources contact: hr@example.invalid | +86 21 5555 0101"]
    path = OUT / "pass" / "employment_letter-pass.docx"
    write_docx(path, "Employment and approved leave confirmation", lines)
    fixtures.append(entry("employment_letter", "pass/" + path.name,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          {"employer_name": "Shanghai Horizon Design Co., Ltd.",
                           "job_title": "Senior Product Designer", "leave_start": "2026-10-02",
                           "leave_end": "2027-01-05", "annual_salary": "CNY 420000"}, [],
                          note="Employer, role, approved leave and return-to-work date are stated."))
    lines = ["To: Entry Clearance Officer", "Employee Name: Mei Ling Chen",
             "Employee ID: SHD-1048", "Employer Name: Shanghai Horizon Design Co., Ltd.",
             "Leave Start: 2026-10-02", "Leave End: 2027-01-05",
             "Annual Salary: CNY 420000",
             "Dear Entry Clearance Officer,",
             "Mei Ling Chen has been granted leave, but this version does not identify her role.",
             "Human Resources contact: hr@example.invalid | +86 21 5555 0101"]
    path = OUT / "fail" / "employment_letter-fail-missing-job-title.docx"
    write_docx(path, "Employment and approved leave confirmation", lines)
    fixtures.append(entry("employment_letter", "fail/" + path.name,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          {"employer_name": "Shanghai Horizon Design Co., Ltd.",
                           "leave_start": "2026-10-02", "leave_end": "2027-01-05",
                           "annual_salary": "CNY 420000"}, ["field_present"], case="fail",
                          note="The letter does not identify the applicant's job title."))

    # Home ties: scanned PDF. The negative is deliberately unreadable.
    lines = ["Property account summary", "Tie Types: apartment mortgage; elderly mother as dependant",
             "Property city: Shanghai", "Mortgage status: active"]
    path = OUT / "pass" / "home_ties_evidence-pass-scanned.pdf"
    render_scan(path, "HOME TIES EVIDENCE", lines, "PDF")
    add_ocr_sidecar(path, lines)
    fixtures.append(entry("home_ties_evidence", "pass/" + path.name, "application/pdf",
                          {"tie_types": "apartment mortgage; elderly mother as dependant"}, [],
                          sidecar=True, note="Scanned property and dependant summary exposes tie types."))
    visual_lines = ["Property account summary", "Image is cropped below the customer details",
                    "Mortgage status: active"]
    path = OUT / "fail" / "home_ties_evidence-fail-cropped-scan.pdf"
    render_scan(path, "HOME TIES EVIDENCE", visual_lines, "PDF")
    add_ocr_sidecar(path, visual_lines)
    fixtures.append(entry("home_ties_evidence", "fail/" + path.name, "application/pdf", {},
                          ["unreadable"], case="fail", sidecar=True,
                          note="Cropped scan contains none of the requested home-tie fields."))

    # Extra paired fixtures exercise cross-document consistency.
    invite_lines = ["Sponsor Name: Wei Zhang",
                    "Sponsor Address: 14 Bramhall Road, Manchester M20 3QT",
                    "Relationship: family friend", "Stay Start: 2026-10-05",
                    "Stay End: 2027-01-03", "Funding Offered: accommodation only"]
    path = OUT / "cross-document" / "sponsor_invitation_letter-fail-name-mismatch.docx"
    write_docx(path, "Invitation letter", invite_lines)
    fixtures.append(entry("sponsor_invitation_letter", "cross-document/" + path.name,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          {"sponsor_name": "Wei Zhang",
                           "sponsor_address": "14 Bramhall Road, Manchester M20 3QT",
                           "relationship": "family friend", "stay_start": "2026-10-05",
                           "stay_end": "2027-01-03", "funding_offered": "accommodation only"},
                          ["cross_document_consistency"], case="fail",
                          note="Invitation sponsor conflicts with the paired status document.",
                          pair="sponsor-name-mismatch"))
    lines = ["Sponsor Name: Hui Chen", "Status Type: Indefinite Leave to Remain"]
    path = OUT / "cross-document" / "sponsor_status_proof-pair-name-mismatch.png"
    render_scan(path, "UK STATUS EVIDENCE", lines, "PNG")
    add_ocr_sidecar(path, lines)
    fixtures.append(entry("sponsor_status_proof", "cross-document/" + path.name, "image/png",
                          {"sponsor_name": "Hui Chen", "status_type": "Indefinite Leave to Remain"},
                          [], sidecar=True, note="Paired status document for mismatch test.",
                          pair="sponsor-name-mismatch"))

    manifest = {
        "fixture_version": "1.0.0",
        "synthetic_data_only": True,
        "applicant": {"applicant_name": "Mei Ling Chen", "trip_start": "2026-10-05",
                      "trip_end": "2027-01-03"},
        "fixtures": fixtures,
    }
    (OUT / "manifest.yaml").write_text(
        yaml.safe_dump(manifest, sort_keys=False, allow_unicode=True), encoding="utf-8")


if __name__ == "__main__":
    generate()
